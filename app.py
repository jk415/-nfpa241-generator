from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
from datetime import datetime

app = Flask(__name__)

RED   = RGBColor(0xDC, 0x26, 0x26)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
GREY  = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_bottom_border(para, color="DC2626", size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_spacing(para, before=0, after=100):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)

def generate_plan(data):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.4)
        section.right_margin  = Inches(1.4)

    def h1(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=240, after=120)
        add_bottom_border(p, "DC2626", 12)
        run = p.add_run(text)
        run.bold = True; run.font.size = Pt(14)
        run.font.color.rgb = DARK; run.font.name = 'Calibri'
        return p

    def h2(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=160, after=80)
        run = p.add_run(text)
        run.bold = True; run.font.size = Pt(12)
        run.font.color.rgb = RED; run.font.name = 'Calibri'
        return p

    def h3(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=120, after=60)
        run = p.add_run(text)
        run.bold = True; run.underline = True
        run.font.size = Pt(11); run.font.color.rgb = DARK; run.font.name = 'Calibri'
        return p

    def body(text, bold=False, italic=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5):
        p = doc.add_paragraph()
        p.alignment = align
        set_para_spacing(p, after=80)
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size); run.font.name = 'Calibri'
        if color: run.font.color.rgb = color
        return p

    def bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        set_para_spacing(p, after=60)
        run = p.add_run(text)
        run.font.size = Pt(10.5); run.font.name = 'Calibri'
        return p

    def kv(key, val):
        p = doc.add_paragraph()
        set_para_spacing(p, after=60)
        r1 = p.add_run(f"{key}:  "); r1.bold = True
        r1.font.size = Pt(10.5); r1.font.name = 'Calibri'
        r2 = p.add_run(val or "[TBD]")
        r2.font.size = Pt(10.5); r2.font.name = 'Calibri'
        return p

    def spacer(after=80):
        p = doc.add_paragraph()
        set_para_spacing(p, after=after)
        return p

    def page_break():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        return p

    def dark_table_row(table, values, header=False):
        row = table.add_row()
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            if header:
                set_cell_bg(cell, "1E293B")
                run = cell.paragraphs[0].add_run(val)
                run.bold = True; run.font.color.rgb = WHITE
            else:
                run = cell.paragraphs[0].add_run(val or "")
                run.font.color.rgb = DARK
            run.font.size = Pt(9.5); run.font.name = 'Calibri'
        return row

    fd       = data.get('fdName') or f"{data.get('city','[CITY]')} Fire Department"
    ahj      = data.get('ahjName') or f"City of {data.get('city','[CITY]')} Inspectional Services Department"
    contacts = data.get('contacts', [])
    fppm_c   = next((c for c in contacts if 'fppm' in c.get('role','').lower() or 'fire prev' in c.get('role','').lower()), contacts[0] if contacts else {})
    fppm_name  = fppm_c.get('name', '[FPPM NAME]')
    fppm_phone = fppm_c.get('phone', '[PHONE]')
    addr     = data.get('projectAddress', '[PROJECT ADDRESS]')
    today    = data.get('planDate', datetime.today().strftime('%B %d, %Y'))
    occupied = data.get('isOccupied', False)

    # ── COVER ──────────────────────────────────────────────────────────────
    spacer(200)
    body("NFPA-241", bold=True, size=32, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("CONSTRUCTION SAFETY PLAN", bold=True, size=26, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(120)
    body(addr, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(80)
    body(f"Date: {today}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    body(data.get('revision', 'Rev. 0 — Initial Submission'), size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(120)
    body("CAP Design Group", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("335 Washington St. Suite 1114, Woburn, MA 01801", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("Jason Kahan P.E. — MA Fire Protection PE License No. 48388", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("617-644-0014  |  admin@capcofire.com  |  capcofire.com", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(120)
    body("[PE SEAL — INSERT HERE]", italic=True, size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break()

    # ── CONTACTS ───────────────────────────────────────────────────────────
    h1("PROJECT CONTACT AND STAKEHOLDER LIST")
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    dark_table_row(t, ["Contact", "Project Role", "Contact Information"], header=True)
    for c in contacts:
        dark_table_row(t, [c.get('name',''), c.get('role',''), c.get('phone','')])
    spacer(160)
    h2("PREPARED FOR:")
    body(ahj); body(fd); body(data.get('fdAddress',''))
    page_break()

    # ── TOC ────────────────────────────────────────────────────────────────
    h1("TABLE OF CONTENTS")
    toc = [
        ("1. ADMINISTRATION / APPLICABLE CODE", True),
        ("   1.1  Plan Overview", False), ("   1.2  Building Summary", False), ("   1.3  Applicable Code & Standards", False),
        ("2. TEMPORARY CONSTRUCTION, EQUIPMENT & STORAGE", True),
        ("   2.1  Construction Fence Requirements", False), ("   2.2  Temporary Structure Construction Requirements", False),
        ("   2.3  Project Site Security Measures", False), ("   2.4  Construction Debris Storage", False),
        ("3. PROCESSES & HAZARDS", True),
        ("   3.1  Fall Protection", False), ("   3.2  PPE Requirements", False), ("   3.3  Hot Work Instructions", False),
        ("   3.4  Smoking", False), ("   3.5  Special Circumstances", False),
        ("4. FIRE PROTECTION", True),
        ("   4.1  Responsibilities", False), ("   4.2  Pre-Incident Planning", False), ("   4.3  Access for Fire Department", False),
        ("   4.4  Emergency Muster Points", False), ("   4.5  Existing Life Safety Systems", False),
        ("   4.6  Site Fire Alarm Notification", False), ("   4.7  Fire Command Post", False),
        ("5. SAFEGUARDING FOR CONSTRUCTION, ALTERATION & DEMOLITION", True),
        ("   5.1  Emergency Evacuation", False), ("   5.2  Emergency Training", False),
    ]
    if occupied:
        toc.append(("   5.3  Occupied Building & Phased Installation Plan", False))
    toc += [
        ("   5.4  Evacuation Plans", False), ("   5.5  Evacuation Signaling & Egress Requirements", False),
        ("ANNEX A — Project Location Map", True),
        ("ANNEX B — Emergency Access / PIP & Fire Hydrant Locations", True),
        ("ANNEX C — Permit Log", True),
        ("ANNEX D — NFPA-241 Weekly Safety Checklist", True),
        ("ANNEX E — Fire Code of Massachusetts Weekly Checklist", True),
    ]
    if occupied:
        toc.append(("ANNEX F — 527 CMR 10.9 Emergency Plan Coordination", True))
    for text, bold in toc:
        p = doc.add_paragraph()
        set_para_spacing(p, after=50)
        run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(10.5); run.font.name = 'Calibri'
        run.font.color.rgb = DARK if bold else GREY
    page_break()

    # ── SECTION 1 ──────────────────────────────────────────────────────────
    h1("1. PLAN ADMINISTRATION / APPLICABLE CODE")
    h2("1.1 OVERVIEW")
    body("This MASTER NFPA-241 PLAN/FIRE PROGRAM has been developed to address the responsibilities required per NFPA-241 for safeguarding during construction and includes site safety requirements to be adhered to throughout construction.")
    body("This NFPA-241 report is intended to be a live document and will require the FPPM, owner, general contractor, and authorities having jurisdiction to provide updates throughout construction in order to remain current.")
    spacer(60)
    body("[ANNEX A — LOCUS MAP TO BE INSERTED HERE]", italic=True, color=GREY)
    body(f"[ANNEX A — NEAREST FIRE DEPARTMENT MAP — {fd} — INSERT HERE]", italic=True, color=GREY)
    body("[ANNEX A — FIRST RESPONDER VANTAGE POINT PHOTOGRAPH — INSERT HERE]", italic=True, color=GREY)
    spacer(60)
    h2("1.2 BUILDING SUMMARY")
    summary = data.get('buildingSummary') or (
        f"The property at {addr} is an existing {data.get('buildingUse','[building use]')} that will "
        f"{'remain operational during construction' if occupied else 'be unoccupied during construction'}. "
        f"The scope of work consists of {data.get('scopeOfWork','[scope of work]')}. No change of use is proposed."
    )
    body(summary)
    if occupied:
        body("The building will remain in active use throughout construction. The building will continue to operate under a 527 CMR 10.9 Emergency Plan, which remains in effect throughout construction. This NFPA 241 Construction Safety Plan supplements the Emergency Plan by addressing construction-specific fire prevention measures, system impairments, occupant notification, and emergency coordination.")
    spacer(60)
    h3("Building Description:")
    kv("Use Group", data.get('useGroup', ''))
    kv("Construction Type", data.get('constructionType', ''))
    kv("Building Height", data.get('buildingHeight', ''))
    kv("Stories Above Grade", data.get('storiesAbove', ''))
    kv("Stories Below Grade", data.get('storiesBelow', ''))
    kv("Occupant Load (Est.)", data.get('occupantLoad', ''))
    kv("Hazardous Materials", data.get('hazardousMaterials', 'None'))
    kv("Emergency Vehicle Access", f"{data.get('accessStreet','[ACCESS STREET]')} — fire lanes {'are existing' if data.get('firelanesExisting') else 'to be established'}")
    spacer(60)
    h2("1.3 APPLICABLE CODE & STANDARDS")
    body("This document is prepared in accordance with the requirements of the following codes and standards:")
    codes = [
        "NFPA-10  (2018)  Standard for Portable Fire Extinguishers",
        "NFPA-13  (2019)  Standard for Installation of Automatic Sprinkler Systems",
        "NFPA-30  (2021)  Flammable and Combustible Liquids",
        "NFPA-51B (2019)  Standard for Fire Prevention During Welding, Cutting, and Hot Work",
        "NFPA-70  (2023)  National Electric Code / 527 CMR 12",
        "NFPA-72  (2019)  National Fire Alarm and Signaling Code",
        "NFPA-101 (2021)  Life Safety Code",
        "NFPA-241 (2019)  Safeguarding During Construction, Alteration & Demolition",
        "780 CMR  (10th Ed.)  Massachusetts State Building Code",
        "527 CMR  (1.00)  Massachusetts Fire Code",
        "521 CMR           Massachusetts Architectural Access Board",
    ]
    if occupied:
        codes.append("527 CMR  (10.9)  Emergency Plan Requirements — Occupied Buildings")
    for c in codes:
        bullet(c)
    page_break()

    # ── SECTION 2 ──────────────────────────────────────────────────────────
    h1("2. TEMPORARY CONSTRUCTION & STORAGE")
    h2("2.1 TEMPORARY FENCING")
    body("All construction fencing is to be minimum 7'-0\" tall chain link with privacy fabric. Gates and openings will be clearly marked with black lettering on white background. Project contact information and building permit shall be posted at the main gate entrance.")
    spacer(60)
    h2("2.2 TEMPORARY SITE STRUCTURES")
    body("In the event a temporary structure is required it shall be separated from the building by a minimum of 20 feet per NFPA 241 Table 4.2.1. All material must be noncombustible or FRT plywood/framing.")
    spacer(40)
    h3("NFPA 241 Table 4.2.1 — Minimum Separation Distances:")
    sep = doc.add_table(rows=1, cols=2)
    sep.style = 'Table Grid'
    dark_table_row(sep, ["Temp Structure Wall Length", "Min. Separation Distance"], header=True)
    for r in [("20 ft (6 m)","30 ft (9 m)"),("30 ft (9 m)","35 ft (11 m)"),("40 ft (12 m)","40 ft (12 m)"),
              ("50 ft (15 m)","45 ft (14 m)"),("60 ft (18 m)","50 ft (15 m)"),(">60 ft","60 ft (18 m)")]:
        dark_table_row(sep, list(r))
    spacer(80)
    h2("2.3 SITE SECURITY MEASURES")
    body("The general contractor is responsible for the integrity of the building throughout construction. The building must be made secure at the end of each working day and inspected for suspicious activity at the start of each working day.")
    spacer(60)
    h2("2.4 CONSTRUCTION DEBRIS STORAGE")
    body("FPPM or their designee will walk the project daily. No substantial quantity of combustible material will be stored inside the building. Dumpsters cannot be left overnight more than 75% filled.")
    page_break()

    # ── SECTION 3 ──────────────────────────────────────────────────────────
    h1("3. PROCESSES & HAZARDS")
    h2("3.1 FALL PROTECTION")
    body("All work at a height of 6'-0\" or greater shall comply with OSHA 29 CFR 1926.502 and NFPA 241 Section 3.1.")
    bullet("No work within 6'-0\" of a leading edge without handrails or personal fall arrest systems")
    bullet("Guardrail systems shall be a minimum of 42\" tall at all vertical openings")
    bullet("All roof work assessed by FPPM prior to commencement — fall arrest systems required")
    bullet("All roof work employees must be OSHA certified for that work type")
    spacer(60)
    h2("3.2 PPE REQUIREMENTS")
    body("All workers are required to wear: approved eye protection, hardhat, appropriate footwear, long pants, shirts, and ear protection. The FPPM is responsible for enforcing PPE requirements.")
    spacer(60)
    h2("3.3 HOT WORK REQUIREMENTS")
    body(f"Any construction activity involving high heat, spark, laser, or flame shall comply with NFPA 51B and the requirements of the {fd}.")
    bullet(f"NO HOT WORK without proper permitting from {fd} and direct sign-off from the FPPM")
    bullet("Any required fire watch shall be provided by the contractor whose work requires it")
    bullet(f"FPPM will coordinate system shutdowns with {fd} prior to any hot work")
    spacer(60)
    h2("3.4 SMOKING")
    body("FPPM will enforce a strict no-smoking policy on and around the property at all times.")
    spacer(60)
    h2("3.5 SPECIAL CIRCUMSTANCES")
    body(data.get('specialCircumstances') or "No special circumstances identified at this time. FPPM shall amend this plan if special circumstances arise during construction.")
    page_break()

    # ── SECTION 4 ──────────────────────────────────────────────────────────
    h1("4. FIRE PROTECTION")
    h2("4.1 RESPONSIBILITIES")
    body(f"The developer/General Contractor is responsible for implementing this NFPA-241 safety plan and designating the FPPM. The FPPM is {fppm_name} ({fppm_phone}).")
    bullet("Walking the project site each morning for security breaches or unsafe material storage")
    bullet("Confirming no large quantities of debris are stored inside or outside the building")
    bullet("Reviewing work areas for conformance with this plan")
    bullet("Maintaining the active permit log, hot works log, and weekly safety checklists (Annex D and E)")
    spacer(60)
    h2("4.2 PRE-INCIDENT PLANNING")
    body(f"FPPM will be familiar with all aspects of the project and provide requested information to {fd} Fire Prevention and {ahj}.")
    bullet("Post emergency contact list next to building permit — visible from exterior")
    bullet(f"At any time conditions change, {fd} shall receive an updated copy without being asked")
    bullet("Approved first aid kits will be provided in a central location")
    spacer(60)
    h2("4.3 FIRE DEPARTMENT SITE ACCESS")
    body(f"Emergency response vehicles will enter the site via {data.get('accessStreet','[ACCESS STREET]')}. Fire lanes are {'existing' if data.get('firelanesExisting') else 'to be established per the site plan'}. Primary entrance is via front door Knox Box.")
    kv("Knox Box Location", data.get('knoxBoxLocation', 'Main entrance'))
    spacer(60)
    h2("4.4 EMERGENCY MUSTER POINTS")
    bullet(f"Primary Muster Point (A): {data.get('musterPointA','[DESCRIBE — See Annex B]')}")
    bullet(f"Secondary Muster Point (B): {data.get('musterPointB','[DESCRIBE — See Annex B]')}")
    spacer(60)
    h2("4.5 EXISTING LIFE SAFETY SYSTEMS")
    body(data.get('existingSystems','During construction the building shall employ fire extinguishers throughout per PIP plans.'))
    spacer(60)
    h2("4.6 SITE FIRE ALARM NOTIFICATION")
    body("Site fire alarm notification shall be via air horn located on site. Upon activation, all workers will proceed to Muster Point A unless instructed otherwise by the FPPM.")
    spacer(60)
    h2("4.7 FIRE COMMAND POST")
    body(f"Located on the ground floor of {addr}. The command post will contain:")
    bullet("Emergency Contact List"); bullet("Master NFPA-241 Plan")
    bullet("General Contractor Project Narrative"); bullet("Set of Current Floor Plans")
    bullet("Key Access (Knox Box keys)"); bullet('"FIRE COMMAND POST" Signage')
    page_break()

    # ── SECTION 5 ──────────────────────────────────────────────────────────
    h1("5. SAFEGUARDING FOR CONSTRUCTION, ALTERATION & DEMOLITION")
    h2("5.1 EMERGENCY PROTOCOL")
    body("In the event of an emergency, all workers will evacuate the building and assemble at one of the muster points. FPPM will meet emergency responders at the site entrance.")
    h3("FIRE EMERGENCY:")
    body("Fire extinguishers throughout site for small incidents. For larger hazards, evacuate and call 911 from outside. No worker shall re-enter until permitted by the FPPM or fire department.")
    h3("MEDICAL EMERGENCY:")
    body("Do not move the victim. Make them comfortable until trained responders arrive.")
    spacer(60)
    h2("5.2 TRAINING")
    body("FPPM shall train all workers on emergency response and preparedness. Each contractor will conduct tool box talks at the start of each work day.")
    if occupied:
        spacer(60)
        h2("5.3 OCCUPIED BUILDING CONSTRUCTION & PHASED INSTALLATION PLAN")
        h3("5.3.1 General")
        body(f"This project involves construction within an existing, occupied {data.get('buildingUse','building')}. The building will remain operational during the installation of {data.get('scopeOfWork','the fire protection and life safety systems')}.")
        h3("5.3.2 Coordination with Occupied Building Operations")
        body("Construction activities shall be coordinated with property management. Occupants shall be provided advance notice (minimum 24 hours; minimum 72 hours for temporary relocation) of activities that may affect access, alarms, noise, or relocations.")
        h3("5.3.3 Phased Construction Approach")
        bullet("Identification of the specific floor or area under construction")
        bullet("Temporary relocation of occupants from the affected floor")
        bullet("Establishment of construction barriers and access controls")
        bullet("Installation of sprinkler and fire alarm rough-in for the designated floor")
        bullet("Restoration of the area and return of occupants upon phase completion")
        if data.get('phaseDescription'): body(data['phaseDescription'])
        h3("5.3.4 Fire Protection and Life Safety During Phased Work")
        body("All required means of egress shall remain available, unobstructed, and clearly marked at all times.")
        h3("5.3.5 Impairments and Compensatory Measures")
        bullet(f"Notification to {fd} when applicable")
        bullet("Implementation of a fire watch when required (see Annex F)")
        bullet("Restoration of systems to service as soon as practicable")
        h3("5.3.6 Hot Work and Housekeeping")
        body("Hot work per NFPA 51B only. Combustible debris removed daily. Fire extinguishers in all active construction areas.")
    spacer(60)
    h2("5.4 EVACUATION PLANS")
    body("The owner and FPPM will coordinate an emergency evacuation plan and post updated egress maps throughout the construction area.")
    spacer(60)
    h2("5.5 EVACUATION SIGNALING & EGRESS REQUIREMENTS")
    body("FPPM shall maintain clear egress passage at all times. All paths shall be clearly lit with emergency lighting throughout construction. Exit signs shall remain illuminated at all times.")
    page_break()

    # ── ANNEX A ────────────────────────────────────────────────────────────
    h1("ANNEX A — PROJECT LOCATION MAP")
    body(f"[LOCUS MAP — {addr} — INSERT SCREENSHOT HERE]", italic=True, color=GREY)
    spacer(300)
    body(f"[NEAREST FIRE DEPARTMENT — {fd} — INSERT MAP SCREENSHOT HERE]", italic=True, color=GREY)
    kv("Nearest Station", data.get('nearestFireStation', f"{fd} — [insert address]"))
    kv("Estimated Response Time", data.get('responseTime', '[confirm with fire department]'))
    spacer(300)
    body("[FIRST RESPONDER VANTAGE POINT — STREET VIEW — INSERT PHOTO HERE]", italic=True, color=GREY)
    page_break()

    # ── ANNEX B ────────────────────────────────────────────────────────────
    h1("ANNEX B — EMERGENCY ACCESS / PRE-INCIDENT PLAN (PIP)")
    body("[FLOOR PLAN — MARK ACCESS POINTS, FIRE HYDRANTS, AND MUSTER POINTS — INSERT HERE]", italic=True, color=GREY)
    spacer(300)
    kv("Primary Access", data.get('accessStreet', '[ACCESS STREET]'))
    kv("Knox Box Location", data.get('knoxBoxLocation', 'Main entrance'))
    kv("Nearest Fire Hydrant", data.get('nearestHydrant', '[confirm with FD]'))
    kv("Muster Point A", data.get('musterPointA', '[front of building]'))
    kv("Muster Point B", data.get('musterPointB', '[rear of property]'))
    spacer(60)
    h2("FIRE HYDRANT LOCATIONS")
    body("[FIRE HYDRANT LOCATION MAP — INSERT GOOGLE MAPS SCREENSHOT ANNOTATED WITH RED ARROWS]", italic=True, color=GREY)
    page_break()

    # ── ANNEX C ────────────────────────────────────────────────────────────
    h1("ANNEX C — PERMIT LOG")
    body("FPPM will update this log throughout construction.")
    spacer(40)
    permit = doc.add_table(rows=1, cols=5)
    permit.style = 'Table Grid'
    dark_table_row(permit, ["Date","Contractor","Description of Work","Hazard Level","Permit No."], header=True)
    for _ in range(8):
        dark_table_row(permit, ["","","","",""])
    page_break()

    # ── ANNEX D ────────────────────────────────────────────────────────────
    h1("ANNEX D — NFPA-241 WEEKLY SAFETY CHECKLIST")
    p = doc.add_paragraph(); set_para_spacing(p, after=80)
    p.add_run("Date: _________________________    Building Permit No.: _________________________").font.size = Pt(10.5)
    p2 = doc.add_paragraph(); set_para_spacing(p2, after=80)
    p2.add_run("Inspector Name: _________________________").font.size = Pt(10.5)
    n241 = doc.add_table(rows=1, cols=3)
    n241.style = 'Table Grid'
    dark_table_row(n241, ["NFPA 241","Description","Y / N / N/A"], header=True)
    for ref, desc in [
        ("4.2","Temporary Offices & Sheds — Separation per Table 4.2.1"),
        ("4.3.4","Fire Extinguishers — Travel distance < 50 ft"),
        ("4.4.1","Equipment — Internal combustion engine exhaust away from combustibles"),
        ("5.1.1","Hot Work — Per NFPA 51B and AHJ policy"),
        ("5.1.3.1","Fire Watch — Dedicated personnel, no other duties assigned"),
        ("5.2.1","Temporary Heating Equipment — Listed and per manufacturer instructions"),
        ("5.2.8","Temporary Heating Equipment — Monitored for safe operation"),
        ("5.4.1","Waste Disposal — Removed daily"),
        ("5.4.4.1","Trash Chute Safety Plan in place"),
        ("5.5.1","Flammable & Combustible Liquids — Max 60 gallons Class I & II"),
        ("5.5.1.5","F/C Liquids & Gases — No Smoking signs posted"),
        ("6.1.1","Electrical — Per NFPA 70 / 527 CMR 12"),
        ("6.1.1.2","Electrical — Extension cords free from damage"),
        ("6.1.2","Temporary Wiring — Branch circuits from approved outlet or panelboard"),
        ("6.1.2.3","Temporary Wiring — Conductors protected by overcurrent devices"),
        ("7.1","Fire Protection — Fire safety program in place"),
        ("7.2.3.1","Fire Protection — Prefire plans developed with fire department"),
        ("7.2.3.3","Fire Protection — On-site FD visit provisions in place"),
        ("7.2.4.4","Weekly self-inspection program active"),
        ("7.4.1","Fire Alarm Reporting — Fire alarm box / telephone service"),
        ("7.5.4","Access for Fire Fighting — Key box in place"),
        ("7.5.6","Stairs — At least one stair provided"),
        ("7.5.6.3","Stairway is lighted"),
        ("7.5.6.5","Exit stairs — Floor level, stair designation, exit direction signage"),
        ("8.6.2.2","Temporary separation walls — One hour fire rating"),
    ]:
        dark_table_row(n241, [ref, desc, ""])
    page_break()

    # ── ANNEX E ────────────────────────────────────────────────────────────
    h1("ANNEX E — FIRE CODE OF MASSACHUSETTS WEEKLY CHECKLIST")
    p = doc.add_paragraph(); set_para_spacing(p, after=80)
    p.add_run("Date: _________________________    Building Permit No.: _________________________").font.size = Pt(10.5)
    p2 = doc.add_paragraph(); set_para_spacing(p2, after=80)
    p2.add_run("Inspector Name: _________________________").font.size = Pt(10.5)
    fc_t = doc.add_table(rows=1, cols=3)
    fc_t.style = 'Table Grid'
    dark_table_row(fc_t, ["IFC / CMR","Description","Y / N / N/A"], header=True)
    for ref, desc in [
        ("1403.1","Temporary Heating Equipment: Listed and Labeled"),
        ("1403.6","Temporary Heating Equipment: Supervised"),
        ("1404.1","Smoking: Approved Areas Designated"),
        ("1404.1;310","Smoking: No Smoking Signs Posted"),
        ("1404.2","Waste Disposal: Combustible Waste Removed Daily"),
        ("1404.5","Fire Watch: When Required by Code Official"),
        ("1404.6;1408.5","Hot Work per AHJ Policy"),
        ("1404.7","Electrical: Temporary Wiring per Building Code Ch. 27"),
        ("1405.1","Flammable & Combustible Liquids: Storage per FC 3404"),
        ("1405.2","F/C Liquids: Class I & II per FC 3406.2"),
        ("1405.4","F/C Liquids: Sources of Ignition / Smoking Prohibited"),
        ("1405.5","F/C Liquids: Class I & II in Safety Containers"),
        ("1406.1","Flammable Gases: Per FC Chapter 35"),
        ("1408.1","Fire Protection: Fire Prevention Superintendent Designated"),
        ("1408.2","Fire Protection: Develop & Maintain Pre-Fire Plan"),
        ("1408.3","Fire Protection: Training of Personnel"),
        ("1408.4","Fire Protection: Fire Protection Devices Maintained"),
        ("1408.6","Fire Protection: Impairments per FC 901"),
        ("1408.7","Fire Protection: Temporary Covering of Devices"),
        ("1409","Fire Alarm Reporting: Emergency Telephone Available"),
        ("1409","Fire Alarm Reporting: Address & Phone Number Posted"),
        ("1410.1","Access for Fire Fighting: Access Roads Available"),
        ("1410.1","Access for Fire Fighting: Key Box Provided"),
        ("1411.1","Means of Egress: One Lighted Stairwell Required"),
        ("1411.2","MOE: Maintained During Construction / Demolition"),
        ("1412.1","Water Supply for Fire Protection: Provided"),
        ("1413","Standpipes"),
        ("1414","Fire Sprinkler: Approved Prior to Occupancy"),
        ("1415","Fire Extinguishers: Each Stairway — All Floors"),
        ("1415","Fire Extinguishers: Each Storage Shed"),
        ("1416.1","Motorized Equipment: Internal Combustion Powered"),
        ("1417","Roofing: Asphalt and Tar Kettles per FC 303"),
        ("1417.3","Roofing: Fire Extinguishers per FC 906"),
    ]:
        dark_table_row(fc_t, [ref, desc, ""])

    # ── ANNEX F (occupied only) ────────────────────────────────────────────
    if occupied:
        page_break()
        h1("ANNEX F — 527 CMR 10.9 EMERGENCY PLAN COORDINATION (OCCUPIED BUILDING)")
        h2("F.1 GENERAL")
        body(f"The building is required to maintain an Emergency Plan in accordance with 527 CMR 10.9. That Emergency Plan remains in effect throughout the duration of construction.")
        spacer(60)
        h2("F.2 LOCATION OF EMERGENCY PLAN")
        bullet(f"Fire Command Post — Ground floor of {addr}")
        bullet("Property Management Office")
        bullet(f"Available to {fd} upon request")
        spacer(60)
        h2("F.3 EMERGENCY CONTACTS")
        fc2 = doc.add_table(rows=1, cols=3)
        fc2.style = 'Table Grid'
        dark_table_row(fc2, ["Name","Role","Phone / Contact"], header=True)
        for c in contacts:
            dark_table_row(fc2, [c.get('name',''), c.get('role',''), c.get('phone','')])
        dark_table_row(fc2, [fd,"AHJ — Fire Prevention","911 / [Confirm with FD]"])
        spacer(80)
        h2("F.4 OCCUPANT NOTIFICATION")
        bullet("Planned construction phases and schedule for each floor")
        bullet("Temporary relocation requirements (minimum 72 hours advance written notice)")
        bullet("Scheduled fire alarm testing or temporary alarm disablement")
        bullet("Any temporary change to egress routes or exits")
        spacer(60)
        h2("F.5 EGRESS AND EVACUATION DURING CONSTRUCTION")
        body("All required means of egress shall remain available, unobstructed, fully lit, and clearly marked at all times.")
        bullet("Upon activation of fire alarm or air horn, all occupants shall evacuate to the nearest available exit")
        bullet("Occupants shall not use elevators — all egress via stairways only")
        bullet(f"Occupants assemble at: {data.get('musterPointA','[OCCUPANT MUSTER POINT]')}")
        spacer(60)
        h2("F.6 SYSTEM IMPAIRMENTS")
        body(f"FPPM shall notify {fd} Fire Prevention before any planned impairment per NFPA 241, NFPA 25, and 527 CMR.")
        spacer(60)
        h2("F.7 FIRE WATCH PROCEDURES (NFPA 241 §5.1.3.1)")
        bullet("Fire watch personnel shall be DEDICATED — no other duties while on watch")
        bullet("Complete patrol of impaired zone every 15 minutes minimum")
        bullet("Each patrol logged: time, areas inspected, inspector name")
        bullet("Watch maintained continuously through impairment including overnight for occupied sleeping areas")
        bullet("Watch not terminated until system fully restored and verified by function test")
        bullet("Personnel shall carry: charged 2A:10B:C extinguisher, communication device, flashlight, fire watch log, and floor plans")
        spacer(60)
        h2("F.8 HOW THE PROPERTY WILL CONTINUE TO OPERATE")
        body(data.get('occupiedOperationNarrative') or "Only one floor or zone shall be in active construction at a time. Occupants on the active construction floor shall be temporarily relocated. Life safety conditions on all non-construction floors shall remain fully intact throughout construction.")
        spacer(60)
        h2(f"F.9 COORDINATION WITH {fd.upper()}")
        bullet(f"FPPM shall meet with {fd} for a pre-construction briefing before any work commences")
        bullet(f"FPPM shall provide {fd} with an updated copy of this plan any time a material change is made")
        bullet(f"{fd} shall be notified at least 24 hours in advance of any planned system impairment or hot work")
        bullet(f"{fd} shall be invited to all final system acceptance tests")

    # ── CERTIFICATION ──────────────────────────────────────────────────────
    page_break()
    h1("CERTIFICATION AND SIGNATURE")
    spacer(60)
    body(f"This NFPA 241 Construction Safety Plan has been prepared for: {addr}")
    spacer(100)
    body("Prepared and Sealed By:", bold=True)
    body("Jason Kahan P.E.")
    body("CAP Design Group")
    body("335 Washington St. Suite 1114, Woburn, MA 01801")
    body("MA Fire Protection PE License No. 48388")
    body("617-644-0014  |  admin@capcofire.com  |  capcofire.com")
    spacer(80)
    body("[PE SEAL — INSERT HERE]", italic=True, color=GREY)
    spacer(100)
    body("_________________________________")
    body("Signature — Jason Kahan P.E.")
    body("Date: _______________________")
    spacer(120)
    body("Fire Protection Program Manager (FPPM):", bold=True)
    body(fppm_name); body(fppm_phone)
    spacer(100)
    body("_________________________________")
    body("FPPM Signature")
    body("Date: _______________________")

    return doc


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()

        doc = generate_plan(data)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        addr_slug = data.get('projectAddress','plan').replace(' ','_').replace(',','')[:40]
        filename  = f"NFPA241_{addr_slug}.docx"

        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return {'error': str(e)}, 500


def generate_battery_calc(data):
    """Generate NFPA 72 Fire Alarm Battery Calculation Document"""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    def h1(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=200, after=100)
        add_bottom_border(p, "DC2626", 12)
        run = p.add_run(text)
        run.bold = True; run.font.size = Pt(14)
        run.font.color.rgb = DARK; run.font.name = 'Calibri'
        return p

    def h2(text):
        p = doc.add_paragraph()
        set_para_spacing(p, before=140, after=60)
        run = p.add_run(text)
        run.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = RED; run.font.name = 'Calibri'
        return p

    def body(text, bold=False, italic=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
        p = doc.add_paragraph()
        p.alignment = align
        set_para_spacing(p, after=60)
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size); run.font.name = 'Calibri'
        if color: run.font.color.rgb = color
        return p

    def kv(key, val):
        p = doc.add_paragraph()
        set_para_spacing(p, after=40)
        r1 = p.add_run(f"{key}:  "); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = 'Calibri'
        r2 = p.add_run(str(val) if val else "[TBD]")
        r2.font.size = Pt(10); r2.font.name = 'Calibri'
        return p

    def spacer(after=60):
        p = doc.add_paragraph()
        set_para_spacing(p, after=after)
        return p

    def calc_row(table, values, header=False, highlight=False):
        row = table.add_row()
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            if header:
                set_cell_bg(cell, "1E293B")
                run = cell.paragraphs[0].add_run(str(val))
                run.bold = True; run.font.color.rgb = WHITE
            elif highlight:
                set_cell_bg(cell, "FEF2F2")
                run = cell.paragraphs[0].add_run(str(val))
                run.bold = True; run.font.color.rgb = RED
            else:
                run = cell.paragraphs[0].add_run(str(val) if val else "")
                run.font.color.rgb = DARK
            run.font.size = Pt(9); run.font.name = 'Calibri'
        return row

    # Extract data
    project_name = data.get('projectName', '[PROJECT NAME]')
    project_addr = data.get('projectAddress', '[PROJECT ADDRESS]')
    facp_model = data.get('facpModel', '[FACP MODEL]')
    calc_date = data.get('calcDate', datetime.today().strftime('%B %d, %Y'))

    standby_hours = float(data.get('standbyHours', 24))
    alarm_minutes = float(data.get('alarmMinutes', 5))
    safety_factor = float(data.get('safetyFactor', 20)) / 100
    derating_factor = float(data.get('deratingFactor', 85)) / 100

    devices = data.get('devices', [])

    # Calculate totals
    total_standby_ma = 0
    total_alarm_ma = 0

    for dev in devices:
        qty = int(dev.get('qty', 0))
        standby = float(dev.get('standbyMa', 0))
        alarm = float(dev.get('alarmMa', 0))
        total_standby_ma += qty * standby
        total_alarm_ma += qty * alarm

    # NFPA 72 Battery Calculation
    # Required capacity = (Standby current x Standby hours) + (Alarm current x Alarm hours)
    # Then apply safety factor and derating
    alarm_hours = alarm_minutes / 60

    standby_ah = (total_standby_ma / 1000) * standby_hours
    alarm_ah = (total_alarm_ma / 1000) * alarm_hours
    subtotal_ah = standby_ah + alarm_ah

    # Apply safety factor
    with_safety = subtotal_ah * (1 + safety_factor)

    # Apply derating factor (batteries should not be discharged below rated capacity)
    required_ah = with_safety / derating_factor

    # ── HEADER ──────────────────────────────────────────────────────────────
    body("FIRE ALARM SYSTEM", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("BATTERY CALCULATION", bold=True, size=18, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("Per NFPA 72 (2019) — National Fire Alarm and Signaling Code", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(40)
    body("CAP Design Group", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("335 Washington St. Suite 1114, Woburn, MA 01801", size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    body("617-644-0014  |  admin@capcofire.com", size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    spacer(80)

    # ── PROJECT INFO ────────────────────────────────────────────────────────
    h1("PROJECT INFORMATION")
    kv("Project Name", project_name)
    kv("Project Address", project_addr)
    kv("FACP Model", facp_model)
    kv("Calculation Date", calc_date)
    kv("Prepared By", data.get('preparedBy', 'CAP Design Group'))
    spacer(40)

    # ── DESIGN CRITERIA ─────────────────────────────────────────────────────
    h1("DESIGN CRITERIA")
    body("Per NFPA 72 Section 10.6.7, secondary power supply shall be capable of operating the fire alarm system under quiescent (standby) load for a minimum of 24 hours, followed by 5 minutes of alarm.")
    spacer(20)
    kv("Standby Duration", f"{standby_hours} hours")
    kv("Alarm Duration", f"{alarm_minutes} minutes")
    kv("Safety Factor", f"{int(safety_factor * 100)}%")
    kv("Battery Derating Factor", f"{int(derating_factor * 100)}%")
    spacer(40)

    # ── DEVICE LOAD TABLE ───────────────────────────────────────────────────
    h1("DEVICE CURRENT LOAD SCHEDULE")

    dev_table = doc.add_table(rows=1, cols=6)
    dev_table.style = 'Table Grid'
    calc_row(dev_table, ["Device Description", "Qty", "Standby (mA)", "Alarm (mA)", "Total Standby (mA)", "Total Alarm (mA)"], header=True)

    for dev in devices:
        qty = int(dev.get('qty', 0))
        standby = float(dev.get('standbyMa', 0))
        alarm = float(dev.get('alarmMa', 0))
        if qty > 0:
            calc_row(dev_table, [
                dev.get('description', ''),
                qty,
                f"{standby:.2f}",
                f"{alarm:.2f}",
                f"{qty * standby:.2f}",
                f"{qty * alarm:.2f}"
            ])

    # Totals row
    calc_row(dev_table, [
        "TOTAL SYSTEM CURRENT",
        "",
        "",
        "",
        f"{total_standby_ma:.2f}",
        f"{total_alarm_ma:.2f}"
    ], highlight=True)

    spacer(60)

    # ── BATTERY CALCULATION ─────────────────────────────────────────────────
    h1("BATTERY CAPACITY CALCULATION")

    calc_table = doc.add_table(rows=1, cols=4)
    calc_table.style = 'Table Grid'
    calc_row(calc_table, ["Description", "Current (A)", "Duration (hrs)", "Capacity (Ah)"], header=True)

    calc_row(calc_table, [
        "Standby Load",
        f"{total_standby_ma / 1000:.3f}",
        f"{standby_hours:.1f}",
        f"{standby_ah:.3f}"
    ])

    calc_row(calc_table, [
        "Alarm Load",
        f"{total_alarm_ma / 1000:.3f}",
        f"{alarm_hours:.4f}",
        f"{alarm_ah:.3f}"
    ])

    calc_row(calc_table, [
        "Subtotal",
        "",
        "",
        f"{subtotal_ah:.3f}"
    ])

    calc_row(calc_table, [
        f"Plus {int(safety_factor * 100)}% Safety Factor",
        "",
        "",
        f"{with_safety:.3f}"
    ])

    calc_row(calc_table, [
        f"Divided by {int(derating_factor * 100)}% Derating",
        "",
        "",
        f"{required_ah:.3f}"
    ], highlight=True)

    spacer(60)

    # ── SUMMARY ─────────────────────────────────────────────────────────────
    h1("CALCULATION SUMMARY")

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'

    calc_row(summary_table, ["Parameter", "Value"], header=True)
    calc_row(summary_table, ["Total Standby Current", f"{total_standby_ma:.2f} mA ({total_standby_ma/1000:.3f} A)"])
    calc_row(summary_table, ["Total Alarm Current", f"{total_alarm_ma:.2f} mA ({total_alarm_ma/1000:.3f} A)"])
    calc_row(summary_table, ["Standby Capacity Required", f"{standby_ah:.3f} Ah"])
    calc_row(summary_table, ["Alarm Capacity Required", f"{alarm_ah:.3f} Ah"])
    calc_row(summary_table, ["MINIMUM BATTERY CAPACITY REQUIRED", f"{required_ah:.2f} Ah"], highlight=True)

    spacer(40)

    # Recommended battery
    std_batteries = [4, 7, 12, 17, 18, 26, 33, 40, 55, 75, 100, 150, 200]
    recommended = next((b for b in std_batteries if b >= required_ah), std_batteries[-1])

    h2("RECOMMENDED BATTERY")
    body(f"Based on the calculation above, the minimum required battery capacity is {required_ah:.2f} Ah.")
    body(f"Recommended Battery: {recommended} Ah (12V sealed lead-acid)", bold=True)

    if recommended >= 26:
        body("Note: For batteries 26 Ah and larger, a separate battery cabinet or enclosure may be required.", italic=True, color=GREY)

    spacer(60)

    # ── NOTES ───────────────────────────────────────────────────────────────
    h1("NOTES AND REFERENCES")
    notes = [
        "1. This calculation is prepared in accordance with NFPA 72 (2019), Section 10.6.7.",
        "2. Standby current values are based on manufacturer specifications at quiescent load.",
        "3. Alarm current values assume all notification appliances operating simultaneously.",
        "4. Battery derating accounts for temperature, age, and capacity reduction over service life.",
        f"5. A {int(safety_factor * 100)}% safety factor has been applied per industry best practice.",
        "6. Actual battery selection should be coordinated with FACP manufacturer requirements.",
        "7. Batteries shall be replaced per manufacturer recommendations or every 5 years, whichever is sooner.",
    ]
    for note in notes:
        body(note, size=9)

    spacer(80)

    # ── CERTIFICATION ───────────────────────────────────────────────────────
    body("_" * 50)
    body("Prepared By:", bold=True)
    body("CAP Design Group")
    body("MA Fire Protection PE License No. 48388")
    spacer(40)
    body("[PE SEAL — INSERT HERE]", italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


@app.route('/battery-calc')
def battery_calc():
    return render_template('battery_calc.html')


@app.route('/generate-battery-calc', methods=['POST'])
def generate_battery_calc_route():
    try:
        data = request.get_json()
        doc = generate_battery_calc(data)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        project_slug = data.get('projectName', 'battery_calc').replace(' ', '_').replace(',', '')[:30]
        filename = f"Battery_Calc_{project_slug}.docx"

        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return {'error': str(e)}, 500


if __name__ == '__main__':
    app.run(debug=True)
